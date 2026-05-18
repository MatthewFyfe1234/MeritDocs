"""
Merit letter template.
Fill in the variables at the top, then run:  python3 merit_letter.py
Output: merit_letter.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# ── FILL THESE IN ──────────────────────────────────────────────────────────────

DATE              = '19 April 2026'
RECIPIENT_NAME    = '<<Recipient Name>>'
COMPANY           = '<<Company>>'
ADDRESS_LINE_1    = '<<Address Line 1>>'
ADDRESS_LINE_2    = '<<Address Line 2>>'
CITY_POSTCODE     = '<<City, Postcode>>'
SUBJECT           = '<<Subject>>'
SALUTATION        = '<<Recipient First Name>>'
BODY_PARAGRAPHS   = [
    '<<Write your message here. Keep paragraphs short and clear.>>',
]
SIGN_OFF          = 'Kind regards,'
SENDER_NAME       = '<<Your Name>>'
SENDER_TITLE      = '<<Your Title>>'
CONFIDENTIALITY   = ''  # optional - leave blank to omit

# ── CONSTANTS ──────────────────────────────────────────────────────────────────

FONT     = 'Inter'
TEXT     = RGBColor(0x49, 0x49, 0x4B)
TEMPLATE = os.path.join(os.path.dirname(__file__), 'Merit_Letter_Template.docx')

# ── DOCUMENT SETUP ─────────────────────────────────────────────────────────────
# Open the template so the header (logo, address, rule) and footer are inherited

doc = Document(TEMPLATE)

# Clear all body content, keeping section properties
body = doc.element.body
for child in list(body):
    if child.tag != qn('w:sectPr'):
        body.remove(child)

# Update header address (right cell of the header table)
for section in doc.sections:
    header = section.header
    for table in header.tables:
        right_cell = table.rows[0].cells[1]
        p = right_cell.paragraphs[0]
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
        address_lines = ['Merit', '20-22 Wenlock Road', 'London, N1 7GU', 'England']
        for i, addr_line in enumerate(address_lines):
            r = p.add_run(addr_line if i == 0 else '\n' + addr_line)
            r.font.name  = FONT
            r.font.size  = Pt(9)
            r.font.color.rgb = TEXT

# Update footer - rebuild main line without phone, handle confidentiality
for section in doc.sections:
    footer = section.footer
    for p in list(footer.paragraphs):
        text = p.text.strip()
        if 'meritoi.com' in text:
            for run in list(p.runs):
                run._element.getparent().remove(run._element)
            r = p.add_run('Merit   \u2022   www.meritoi.com   \u2022   enquiries@meritoi.com')
            r.font.name  = FONT
            r.font.size  = Pt(9)
            r.font.color.rgb = TEXT
        elif 'Optional' in text:
            if CONFIDENTIALITY:
                for run in p.runs:
                    run.text = ''
                r = p.add_run(CONFIDENTIALITY)
                r.font.name = FONT
                r.font.size = Pt(8)
                r.font.color.rgb = TEXT
            else:
                p._element.getparent().remove(p._element)


# ── HELPERS ────────────────────────────────────────────────────────────────────

def line(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name  = FONT
    r.font.bold  = bold
    r.font.color.rgb = TEXT
    return p

def blank():
    line('')

def address_block(lines):
    p = doc.add_paragraph()
    for i, text in enumerate(lines):
        r = p.add_run(('' if i == 0 else '\n') + text)
        r.font.name  = FONT
        r.font.color.rgb = TEXT


# ── CONTENT ────────────────────────────────────────────────────────────────────

blank()
line(DATE, align=WD_ALIGN_PARAGRAPH.RIGHT)
blank()
address_block([RECIPIENT_NAME, COMPANY, ADDRESS_LINE_1, ADDRESS_LINE_2, CITY_POSTCODE])
blank()
line(f'Subject: {SUBJECT}', bold=True)
blank()
line(f'Dear {SALUTATION},')
blank()
for para in BODY_PARAGRAPHS:
    line(para)
    blank()
line(SIGN_OFF)
blank()
address_block([SENDER_NAME, SENDER_TITLE, 'Merit'])


# ── OUTPUT ─────────────────────────────────────────────────────────────────────

out = os.path.join(os.path.dirname(__file__), 'merit_letter.docx')
doc.save(out)
print(f'Written: {out}')
