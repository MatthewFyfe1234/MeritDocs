"""
Merit capabilities document.
Run with:  python3 merit_capabilities.py
Output: merit_capabilities.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import os

# Brand colours
PRIMARY  = RGBColor(0x33, 0x5D, 0x57)
H2_COLOR = RGBColor(0x4A, 0x7E, 0x78)
TEXT     = RGBColor(0x49, 0x49, 0x4B)
FONT     = 'Inter'
TEMPLATE = os.path.join(os.path.dirname(__file__), 'Merit_Letter_Template.docx')

# Open template to inherit header and footer
doc = Document(TEMPLATE)

body = doc.element.body
for child in list(body):
    if child.tag != qn('w:sectPr'):
        body.remove(child)

# Update header address (right cell of the header table)
for section in doc.sections:
    header = section.header
    for table in header.tables:
        right_cell = table.rows[0].cells[1]
        p = right_cell.paragraphs[0]
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
        address_lines = ['Merit', '20-22 Wenlock Road', 'London, N1 7GU', 'England']
        for i, addr_line in enumerate(address_lines):
            r = p.add_run(addr_line if i == 0 else '\n' + addr_line)
            r.font.name  = FONT
            r.font.size  = Pt(9)
            r.font.color.rgb = TEXT

# Clean up footer - rebuild main line without phone, remove optional/blank paragraphs
for section in doc.sections:
    section.footer_distance = Pt(6)
    footer = section.footer
    for p in list(footer.paragraphs):
        text = p.text.strip()
        if not text or 'Optional' in text:
            p._element.getparent().remove(p._element)
        elif 'meritoi.com' in text:
            for run in list(p.runs):
                run._element.getparent().remove(run._element)
            r = p.add_run('Merit   \u2022   www.meritoi.com   \u2022   enquiries@meritoi.com')
            r.font.name  = FONT
            r.font.size  = Pt(9)
            r.font.color.rgb = TEXT


# ── HELPERS ────────────────────────────────────────────────────────────────────

def styled_run(paragraph, text, bold=False, color=TEXT, size=Pt(11)):
    run = paragraph.add_run(text)
    run.font.name  = FONT
    run.font.size  = size
    run.font.bold  = bold
    run.font.color.rgb = color
    return run

def title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    styled_run(p, text, bold=True, color=PRIMARY, size=Pt(20))

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    styled_run(p, text, bold=True, color=PRIMARY, size=Pt(14))

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    styled_run(p, text, bold=True, color=H2_COLOR, size=Pt(11.5))

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    styled_run(p, text)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    styled_run(p, text)


# ── TITLE ──────────────────────────────────────────────────────────────────────

title('Merit: What We Can Do for You')

body(
    'Merit is an operational platform built for organisations that deliver complex products. '
    'It automates the planning, scheduling, and management functions that are common to every '
    'project and every department - consolidating what are typically a collection of disconnected '
    'tools into one connected system.'
)
body('Across your business, Merit delivers three things:')
bullet('Consolidation - replacing fragmented spreadsheets, systems, and manual processes with a single platform')
bullet('Unity - giving every department a shared, live view of the organisation and its performance')
bullet('Transparency - making execution, contribution, and improvement visible at every level')


# ── PROJECT DELIVERY ───────────────────────────────────────────────────────────

h1('Project Delivery')
body(
    'Merit structures the full project lifecycle - from first enquiry through to delivery - '
    'around your products and processes. Once a specification is defined, the system derives '
    'everything downstream automatically: time, materials, scheduling, and resource allocation. '
    'These models are rules-based and self-improving, meaning they get more accurate over time '
    'as real delivery data is recorded against them.'
)

h2('Estimation & Tender')
bullet('Constraint-based product configurator ensures every specification is valid and complete before it leaves the desk - reducing errors and dramatically speeding up the estimating process')
bullet('Time forecasting derived directly from the spec and your predefined process model - not guesswork')
bullet('Material and resource usage calculated automatically from the spec, using predefined ratios that account for waste and are refined over time as actual consumption is recorded')
bullet('Any measurable project output - time, materials, carbon, electricity, cost - can be projected from the same model')
bullet('Delivery timeline generated from forecast data and integrated team calendars, accounting for all commitments, not just project assignments')

h2('Project Planning & Scheduling')
bullet('On acceptance, a full project roadmap is automatically generated from your process model - no manual planning required')
bullet('A complete Gantt plan is produced, scheduling every task across the project timeline')
bullet('Stage gating enforced through document uploads - projects cannot progress past defined stages without required documentation in place')
bullet('Plans adapt in real time as conditions change, with live visibility across all active projects')

h2('Resource & Executor Selection')
bullet('Tasks are assigned to the most suitable person or team based on skills, experience, and live calendar availability')
bullet('Equipment and tooling can carry calendars just like people - so resource allocation accounts for machinery availability as well as workforce capacity')
bullet('Historical performance data informs future assignments, so the best operators and tools are identified and utilised')

h2('Materials & Inventory')
bullet('Inventory requests are generated automatically from the project spec and process model at the point of acceptance')
bullet('Required materials are calculated and flagged against current stock levels, triggering procurement where needed')
bullet('Actual material usage recorded at delivery feeds back into the forecasting model, continuously improving future estimates')


# ── OPERATIONAL FUNCTIONS ──────────────────────────────────────────────────────

h1('Operational Functions')
body(
    'Across your support functions, Merit aggregates data from all project activity into business '
    'intelligence that surfaces performance, identifies improvement, and drives better decisions '
    'across the organisation.'
)

h2('Operations & Business Intelligence')
bullet('Process-integrated BI identifies bottlenecks, inefficiencies, and opportunities for improvement - grounded in real execution data, not periodic reports')
bullet('Identify your best operators, tools, and materials through the data - and use those insights to drive standards across the organisation')
bullet('Compare performance across teams, time periods, clients, materials, processes, tools, and more - giving you a clear picture of where your business excels and where it can improve')
bullet('BI is embedded in the flow of work, so insight is always current and always traceable to its source')

h2('HR')
bullet('Employee onboarding managed through Merit from day one - every new starter has access to the processes they will follow, their performance benchmarks, and all onboarding materials')
bullet('Document management for all employee records - identification, training certificates, and professional certifications can be required, uploaded, and tracked centrally')
bullet('Performance reviewed continuously against peers with equivalent skills and experience - giving a fair, data-driven basis for salary decisions and recognition')
bullet('Workforce planning informed by live redundancy data - understand where your teams are over or under-resourced and plan hiring accordingly')
bullet('Full holiday and absence management with configurable allowance rules - requests reviewed automatically against live capacity data')

h2('Inventory & Stock Control')
bullet('Tools and equipment managed with their own calendars - giving you full visibility of utilisation and availability across your asset base')
bullet('Tool procurement planning driven by live redundancy data - know when to buy before a shortage becomes a problem')
bullet('Stock levels and locations tracked in real time - always know how much you have and where it is')
bullet('Inventory ordering automated through a generatively produced RFQ system - reducing manual procurement effort and improving consistency')
bullet('A full suite of manual inventory tools also available for cases requiring direct intervention')


# ── OUTPUT ─────────────────────────────────────────────────────────────────────

out = os.path.join(os.path.dirname(__file__), 'merit_capabilities.docx')
doc.save(out)
print(f'Written: {out}')
